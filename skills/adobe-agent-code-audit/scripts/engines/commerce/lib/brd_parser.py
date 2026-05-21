"""BRD Parser
=================
Parses BRD files (.txt, .docx, .json) into the structured dict
used by BRDAnalysisEngine.
"""

import json
import os
import re

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def parse_brd_file(filepath):
    """Parse a BRD file (.txt, .docx, or .json) and return structured dict.

    Returns dict with keys: metadata, requirements, technical_context, etc.
    """
    if not filepath or not os.path.isfile(filepath):
        return None

    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".json":
        return _parse_json(filepath)
    elif ext in (".docx", ".doc"):
        return _parse_docx(filepath)
    else:
        return _parse_txt(filepath)


def _parse_docx(filepath):
    """Extract structured BRD data directly from .docx paragraphs + tables."""
    if not HAS_DOCX:
        print(f"   ❌ python-docx is required to parse .docx files. Install: pip install python-docx")
        return None
    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        print(f"   ❌ Failed to open .docx BRD: {e}")
        return None

    paragraphs = [p.text for p in doc.paragraphs]
    tables = doc.tables

    # ── Metadata from Table 1 (key-value pairs) ─────────────────
    metadata = _docx_parse_metadata(tables)

    # ── Requirements from paragraphs + per-REQ affected-area tables ──
    requirements = _docx_parse_requirements(paragraphs, tables)

    # ── Business Rules ───────────────────────────────────────────
    business_rules = _docx_parse_kv_table(tables, "Rule ID", "Rule",
                                           id_key="id", val_key="rule")

    # ── API Contracts ────────────────────────────────────────────
    api_contracts = _docx_parse_api_contracts(tables, paragraphs)

    # ── Data Model ───────────────────────────────────────────────
    data_model = _docx_parse_data_model(tables)

    # ── Technical Context ────────────────────────────────────────
    technical_context = _docx_parse_kv_table_as_dict(tables, "Current Version")

    # ── NFR ──────────────────────────────────────────────────────
    nfr = _docx_parse_kv_table(tables, "Category", "Requirement",
                                id_key="category", val_key="requirement")

    # ── Test Scenarios ───────────────────────────────────────────
    test_scenarios = _docx_parse_kv_table(tables, "Test Case", "Scenario",
                                           id_key="id", val_key="scenario")

    # ── Risks ────────────────────────────────────────────────────
    risks = _docx_parse_kv_table(tables, "Risk / Constraint", "Mitigation",
                                  id_key="risk", val_key="mitigation")

    # ── Patch / Bug (from paragraphs — usually "Not applicable") ─
    patch_details = _docx_parse_patch_section(paragraphs)
    bug_details = {"bugs": []}

    return {
        "metadata": metadata,
        "requirements": requirements,
        "business_rules": business_rules,
        "api_contracts": api_contracts,
        "data_model": data_model,
        "technical_context": technical_context,
        "nfr": nfr,
        "test_scenarios": test_scenarios,
        "risks": risks,
        "patch_details": patch_details,
        "bug_details": bug_details,
    }


# ── Docx helper functions ───────────────────────────────────────────


def _docx_table_to_rows(table):
    """Convert a docx table to list of row dicts (header-keyed)."""
    rows_data = []
    headers = [c.text.strip() for c in table.rows[0].cells]
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        rows_data.append(dict(zip(headers, cells)))
    return headers, rows_data


def _docx_find_table(tables, first_cell_text):
    """Find a table whose first header cell matches the given text."""
    for t in tables:
        if t.rows and t.rows[0].cells:
            if t.rows[0].cells[0].text.strip() == first_cell_text:
                return t
    return None


def _docx_parse_metadata(tables):
    """Parse metadata from the first key-value table (Title / Version / Author)."""
    meta = {
        "title": "", "version": "", "author": "", "date": "",
        "type": "new_requirement", "priority": "high",
        "adobe_commerce_version": "", "architecture": "",
        "modules_affected": [], "tags": [],
    }
    if not tables:
        return meta

    # First table is usually the document info table
    t = tables[0]
    for row in t.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2:
            key = cells[0].lower()
            val = cells[1]
            if "title" in key:
                meta["title"] = val
            elif key == "version":
                meta["version"] = val
            elif "author" in key:
                meta["author"] = val
            elif "date" in key:
                meta["date"] = val
            elif "type" in key or "analysis" in key:
                meta["type"] = val.lower().replace(" ", "_").replace("/", "_")
            elif "priority" in key:
                meta["priority"] = val.lower()
            elif "adobe commerce" in key or "magento" in key:
                meta["adobe_commerce_version"] = val
            elif "architecture" in key:
                meta["architecture"] = val
            elif "module" in key:
                meta["modules_affected"] = [m.strip() for m in val.split(",") if m.strip()]
            elif "tag" in key:
                meta["tags"] = [t.strip() for t in val.split(",") if t.strip()]
    return meta


def _docx_parse_requirements(paragraphs, tables):
    """Parse requirements from paragraphs (REQ-NNN blocks) and affected area tables."""
    requirements = []
    # Find all REQ-NNN blocks in paragraphs
    req_indices = []
    for i, line in enumerate(paragraphs):
        m = re.match(r'(REQ-\d+)\s*:\s*(.+)', line.strip())
        if m:
            req_indices.append((i, m.group(1), m.group(2).strip()))

    # Track which table index to use for affected areas (Tables 2..N are per-REQ)
    # Find tables that start with "Modules" in first cell (affected area tables)
    affected_tables = []
    for t in tables:
        if t.rows and t.rows[0].cells:
            first = t.rows[0].cells[0].text.strip()
            if first == "Modules":
                affected_tables.append(t)

    for idx, (para_idx, req_id, req_title) in enumerate(req_indices):
        # Collect content until next REQ or next numbered section
        lines_after = []
        for j in range(para_idx + 1, len(paragraphs)):
            line = paragraphs[j].strip()
            # Stop at next REQ-NNN or next top-level section
            if re.match(r'REQ-\d+\s*:', line):
                break
            if re.match(r'^\d+\.\s+[A-Z]', line) and not line.startswith(req_id):
                break
            lines_after.append(line)

        # Extract description
        description = ""
        ac_lines = []
        deps = []
        in_desc = False
        in_ac = False
        in_deps = False

        for line in lines_after:
            if line.startswith("Description:"):
                in_desc = True
                in_ac = False
                in_deps = False
                desc_part = line[len("Description:"):].strip()
                if desc_part:
                    description = desc_part
                continue
            elif line.startswith("Acceptance Criteria:"):
                in_desc = False
                in_ac = True
                in_deps = False
                continue
            elif line.startswith("Affected Areas:"):
                in_desc = False
                in_ac = False
                in_deps = False
                continue
            elif line.startswith("Dependencies:"):
                in_desc = False
                in_ac = False
                in_deps = True
                continue

            if in_desc and line:
                description += " " + line if description else line
            elif in_ac and line:
                if re.match(r'AC\d+:', line):
                    ac_lines.append(line)
            elif in_deps and line:
                deps.append(line.lstrip("- "))

        # Get affected areas from corresponding table
        affected = {"modules": [], "flows": [], "apis": [], "events": [], "tables": [], "admin_pages": []}
        if idx < len(affected_tables):
            at = affected_tables[idx]
            for row in at.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    key = cells[0].lower()
                    vals = [v.strip() for v in cells[1].split(",") if v.strip()]
                    if "module" in key:
                        affected["modules"] = vals
                    elif "flow" in key:
                        affected["flows"] = vals
                    elif "api" in key:
                        affected["apis"] = vals
                    elif "event" in key:
                        affected["events"] = vals
                    elif "table" in key or "db" in key:
                        affected["tables"] = vals
                    elif "admin" in key:
                        affected["admin_pages"] = vals

        requirements.append({
            "id": req_id,
            "title": req_title,
            "description": description,
            "acceptance_criteria": ac_lines,
            "affected_areas": affected,
            "dependencies": deps,
        })

    return requirements


def _docx_parse_kv_table(tables, key_header, val_header, id_key="id", val_key="value"):
    """Parse a 2-column key-value table into list of dicts."""
    for t in tables:
        if not t.rows:
            continue
        headers = [c.text.strip() for c in t.rows[0].cells]
        if len(headers) >= 2 and headers[0] == key_header and headers[1] == val_header:
            items = []
            for row in t.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0]:
                    items.append({id_key: cells[0], val_key: cells[1]})
            return items
    return []


def _docx_parse_api_contracts(tables, paragraphs):
    """Parse API contract tables (queries, mutations) and error codes from paragraphs."""
    result = {"queries": [], "mutations": [], "error_codes": []}

    for t in tables:
        if not t.rows:
            continue
        headers = [c.text.strip() for c in t.rows[0].cells]
        if len(headers) >= 2:
            if headers[0] == "Query" and headers[1] == "Purpose":
                for row in t.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) >= 2 and cells[0]:
                        result["queries"].append({"name": cells[0], "purpose": cells[1]})
            elif headers[0] == "Mutation" and headers[1] == "Purpose":
                for row in t.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) >= 2 and cells[0]:
                        result["mutations"].append({"name": cells[0], "purpose": cells[1]})

    # Error codes from paragraphs (uppercase lines after "Recommended Error Codes")
    in_codes = False
    for line in paragraphs:
        t = line.strip()
        if "Error Codes" in t:
            in_codes = True
            continue
        if in_codes:
            if t and t == t.upper() and "_" in t:
                result["error_codes"].append(t)
            elif t and not t.upper() == t:
                in_codes = False

    return result


def _docx_parse_data_model(tables):
    """Parse data model from 3-column table (Table / Entity, Key Fields, Purpose)."""
    for t in tables:
        if not t.rows:
            continue
        headers = [c.text.strip() for c in t.rows[0].cells]
        if len(headers) >= 3 and "Table" in headers[0]:
            items = []
            for row in t.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 3 and cells[0]:
                    items.append({
                        "table": cells[0],
                        "fields": cells[1],
                        "purpose": cells[2],
                    })
            return items
    return []


def _docx_parse_kv_table_as_dict(tables, first_key):
    """Parse a 2-column key-value table into a flat dict (for technical context)."""
    result = {}
    for t in tables:
        if not t.rows:
            continue
        first_cell = t.rows[0].cells[0].text.strip()
        if first_cell == first_key:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0]:
                    key = cells[0].lower().replace(" ", "_").replace("/", "_")
                    result[key] = cells[1]
            return result
    return result


def _docx_parse_patch_section(paragraphs):
    """Parse patch details from paragraphs (usually 'Not applicable')."""
    default = {"from_version": "", "to_version": "", "patch_ids": [],
                "deprecated_classes": [], "removed_methods": [],
                "changed_interfaces": [], "db_schema_changes": []}

    in_patch = False
    for line in paragraphs:
        t = line.strip()
        if re.match(r'\d+\.\s*Patch\s*/?\s*Upgrade', t, re.IGNORECASE):
            in_patch = True
            continue
        if in_patch:
            if "not applicable" in t.lower():
                return default
            if re.match(r'\d+\.\s+', t):
                break
    return default


def _parse_json(filepath):
    """Parse JSON BRD file directly."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (json.JSONDecodeError, IOError) as e:
        print(f"   ❌ Failed to parse BRD JSON: {e}")
        return None


def _parse_txt(filepath):
    """Parse the structured .txt BRD format into a dict."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        print(f"   ❌ BRD file is not a valid text file (binary?): {filepath}")
        print(f"      Use .docx format instead, or convert to .txt.")
        return None
    except IOError as e:
        print(f"   ❌ Failed to read BRD file: {e}")
        return None
    return _parse_content(content)


def _parse_content(content):
    """Parse BRD text content (from .txt or extracted from .docx) into structured dict."""

    result = {
        "metadata": _parse_metadata(content),
        "requirements": _parse_requirements(content),
        "business_rules": _parse_business_rules(content),
        "api_contracts": _parse_api_contracts(content),
        "data_model": _parse_data_model(content),
        "technical_context": _parse_technical_context(content),
        "nfr": _parse_nfr(content),
        "test_scenarios": _parse_test_scenarios(content),
        "patch_details": _parse_patch_details(content),
        "bug_details": _parse_bug_details(content),
    }

    return result


def _extract_section(content, section_marker):
    """Extract content between a section header and the next section or EOF.
    Matches flexibly: tries exact match first, then keyword-based match
    to handle varying section numbers across BRD templates.
    """
    escaped = re.escape(section_marker)
    # Try exact match first
    pattern = r'={10,}\n' + escaped + r'\n={10,}'
    match = re.search(pattern, content)

    # If exact match fails, try flexible match by section keyword
    if not match:
        # Extract the meaningful part after "SECTION N: "
        keyword = re.sub(r'^SECTION\s+\d+:\s*', '', section_marker)
        if keyword != section_marker:
            # Search for any "SECTION X: <keyword>" pattern
            flex_pattern = r'={10,}\nSECTION\s+\d+:\s*' + re.escape(keyword) + r'\n={10,}'
            match = re.search(flex_pattern, content)

    if not match:
        return ""

    start = match.end()

    # Find next section (next ===...=== block)
    next_section = re.search(r'\n={10,}\n', content[start:])
    if next_section:
        end = start + next_section.start()
    else:
        end = len(content)

    return content[start:end].strip()


def _extract_field(text, field_name):
    """Extract a 'Field : value' field from text."""
    pattern = rf'{re.escape(field_name)}\s*:\s*(.+?)(?:\n|$)'
    match = re.search(pattern, text, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return ""


def _split_csv(value):
    """Split a comma-separated value string into a list."""
    if not value:
        return []
    items = [item.strip() for item in value.split(",")]
    return [item for item in items if item and not item.startswith("[")]


def _parse_metadata(content):
    """Parse Document Information section."""
    # Look for the document info block between first ═ lines or ─ lines
    doc_info = ""
    info_match = re.search(r'Document Information\s*\n[─]+\n(.*?)(?:\n\n={10,}|\n={10,})', content, re.DOTALL)
    if info_match:
        doc_info = info_match.group(1)
    else:
        # Fallback: look between first ═ block and SECTION 1
        first_block = re.search(r'={10,}\n\n(.*?)={10,}\s*\nSECTION 1', content, re.DOTALL)
        if first_block:
            doc_info = first_block.group(1)

    title = _extract_field(doc_info, "Title")
    version = _extract_field(doc_info, "Version")
    author = _extract_field(doc_info, "Author")
    date = _extract_field(doc_info, "Date")
    analysis_type = _extract_field(doc_info, "Analysis Type")
    priority = _extract_field(doc_info, "Priority")
    adobe_commerce = _extract_field(doc_info, "Adobe Commerce")
    architecture = _extract_field(doc_info, "Architecture")
    modules_affected = _split_csv(_extract_field(doc_info, "Modules Affected"))
    tags = _split_csv(_extract_field(doc_info, "Tags"))

    return {
        "title": title,
        "version": version,
        "author": author,
        "date": date,
        "type": analysis_type.lower().replace(" ", "_").replace("/", "_"),
        "priority": priority.lower(),
        "adobe_commerce_version": adobe_commerce,
        "architecture": architecture,
        "modules_affected": modules_affected,
        "tags": tags,
    }


def _parse_requirements(content):
    """Parse SECTION N: REQUIREMENTS into a list of requirement dicts."""
    section = _extract_section(content, "SECTION 1: REQUIREMENTS")
    if not section:
        # Try alternate numbering
        section = _extract_section(content, "SECTION 2: REQUIREMENTS")
    if not section:
        return []

    # Split by requirement blocks (─── lines followed by Requirement ID)
    req_blocks = re.split(r'[─]{10,}\s*\nRequirement ID\s*:', section)
    requirements = []

    for block in req_blocks[1:]:  # Skip the intro text before first requirement
        block = "Requirement ID  :" + block  # Restore the split part
        req = _parse_single_requirement(block)
        if req:
            requirements.append(req)

    return requirements


def _parse_single_requirement(block):
    """Parse a single requirement block."""
    req_id = _extract_field(block, "Requirement ID")
    title = _extract_field(block, "Title")

    # Parse description (multiline after "Description:")
    desc = _extract_multiline_field(block, "Description")

    # Parse acceptance criteria
    ac_lines = []
    ac_match = re.search(r'Acceptance Criteria:\s*\n(.*?)(?:\nAffected Areas:|\nDependencies:|\Z)', block, re.DOTALL)
    if ac_match:
        for line in ac_match.group(1).strip().split('\n'):
            line = line.strip()
            if line and re.match(r'AC\d+:', line):
                ac_lines.append(line)

    # Parse affected areas
    affected = _parse_affected_areas(block)

    # Parse dependencies
    deps = []
    dep_match = re.search(r'Dependencies:\s*\n(.*?)(?:\n[─]{10,}|\n={10,}|\Z)', block, re.DOTALL)
    if dep_match:
        for line in dep_match.group(1).strip().split('\n'):
            line = line.strip().lstrip('- ')
            if line and not line.startswith('['):
                deps.append(line)

    return {
        "id": req_id,
        "title": title,
        "description": desc,
        "acceptance_criteria": ac_lines,
        "affected_areas": affected,
        "dependencies": deps,
    }


def _extract_multiline_field(block, field_name):
    """Extract a multiline field (indented lines after 'Field:')."""
    pattern = rf'{re.escape(field_name)}:\s*\n(.*?)(?:\n\w|\nAcceptance|\nAffected|\nDependencies|\Z)'
    match = re.search(pattern, block, re.DOTALL)
    if match:
        lines = match.group(1).strip().split('\n')
        return ' '.join(line.strip() for line in lines if line.strip())
    return ""


def _parse_affected_areas(block):
    """Parse the Affected Areas sub-section."""
    areas_match = re.search(r'Affected Areas:\s*\n(.*?)(?:\nDependencies:|\n[─]{10,}|\n={10,}|\Z)', block, re.DOTALL)
    if not areas_match:
        return {"modules": [], "flows": [], "apis": [], "events": [], "tables": [], "admin_pages": []}

    areas_text = areas_match.group(1)
    return {
        "modules": _split_csv(_extract_field(areas_text, "Modules")),
        "flows": _split_csv(_extract_field(areas_text, "Flows")),
        "apis": _split_csv(_extract_field(areas_text, "APIs")),
        "events": _split_csv(_extract_field(areas_text, "Events")),
        "tables": _split_csv(_extract_field(areas_text, "DB Tables")),
        "admin_pages": _split_csv(_extract_field(areas_text, "Admin Pages")),
    }


def _parse_technical_context(content):
    """Parse SECTION N: TECHNICAL CONTEXT."""
    section = _extract_section(content, "SECTION 2: TECHNICAL CONTEXT")
    if not section:
        section = _extract_section(content, "SECTION 6: TECHNICAL CONTEXT")
    if not section:
        return {}

    return {
        "current_version": _extract_field(section, "Current Version"),
        "target_version": _extract_field(section, "Target Version"),
        "php_version": _extract_field(section, "PHP Version"),
        "frontend": _extract_field(section, "Frontend"),
        "commerce_apis": _extract_field(section, "Commerce APIs"),
        "custom_modules": _split_csv(_extract_field(section, "Custom Modules")),
        "integrations": _split_csv(_extract_field(section, "Integrations")),
        "multi_store": _extract_field(section, "Multi-store"),
        "currency": _extract_field(section, "Currency"),
        "idempotency": _extract_field(section, "Idempotency"),
        "caching": _extract_field(section, "Caching"),
        "performance": _extract_field(section, "Performance"),
        "notes": _extract_field(section, "Notes"),
    }


def _parse_patch_details(content):
    """Parse SECTION N: PATCH / UPGRADE DETAILS."""
    section = _extract_section(content, "SECTION 3: PATCH / UPGRADE DETAILS")
    if not section:
        section = _extract_section(content, "SECTION 12: PATCH / UPGRADE DETAILS")
    if not section:
        return {"from_version": "", "to_version": "", "patch_ids": [],
                "deprecated_classes": [], "removed_methods": [],
                "changed_interfaces": [], "db_schema_changes": []}

    from_ver = _extract_field(section, "From Version")
    to_ver = _extract_field(section, "To Version")
    patch_ids = _split_csv(_extract_field(section, "Patch IDs"))

    deprecated = _parse_numbered_list(section, "Deprecated Classes")
    removed = _parse_numbered_list(section, "Removed Methods")
    changed = _parse_numbered_list(section, "Changed Interfaces")
    db_changes = _parse_numbered_list(section, "DB Schema Changes")

    return {
        "from_version": from_ver,
        "to_version": to_ver,
        "patch_ids": patch_ids,
        "deprecated_classes": deprecated,
        "removed_methods": removed,
        "changed_interfaces": changed,
        "db_schema_changes": db_changes,
    }


def _parse_numbered_list(section, header):
    """Parse a numbered list section (1. item, 2. item, ...)."""
    pattern = rf'{re.escape(header)}[^:]*:\s*\n(.*?)(?:\n\w[^\d]|\n\n[A-Z]|\Z)'
    match = re.search(pattern, section, re.DOTALL)
    if not match:
        return []

    items = []
    for line in match.group(1).strip().split('\n'):
        line = line.strip()
        # Remove numbering prefix like "1. ", "2. "
        cleaned = re.sub(r'^\d+\.\s*', '', line)
        if cleaned and not cleaned.startswith('['):
            items.append(cleaned)
    return items


def _parse_bug_details(content):
    """Parse SECTION N: BUG DETAILS."""
    section = _extract_section(content, "SECTION 4: BUG DETAILS")
    if not section:
        section = _extract_section(content, "SECTION 13: BUG DETAILS")
    if not section or "Not applicable" in section:
        return {"bugs": []}

    # Split by bug blocks
    bug_blocks = re.split(r'[─]{10,}\s*\nBug ID\s*:', section)
    bugs = []

    for block in bug_blocks[1:]:
        block = "Bug ID          :" + block
        bug = _parse_single_bug(block)
        if bug:
            bugs.append(bug)

    return {"bugs": bugs}


def _parse_single_bug(block):
    """Parse a single bug block."""
    bug_id = _extract_field(block, "Bug ID")
    title = _extract_field(block, "Title")
    severity = _extract_field(block, "Severity").lower()
    reported_by = _extract_field(block, "Reported By")
    environment = _extract_field(block, "Environment")

    description = _extract_multiline_field(block, "Description")
    expected = _extract_multiline_field(block, "Expected Behavior")
    actual = _extract_multiline_field(block, "Actual Behavior")

    # Steps to reproduce
    steps = []
    steps_match = re.search(r'Steps to Reproduce:\s*\n(.*?)(?:\nExpected|\nActual|\Z)', block, re.DOTALL)
    if steps_match:
        for line in steps_match.group(1).strip().split('\n'):
            line = line.strip()
            cleaned = re.sub(r'^\d+\.\s*', '', line)
            if cleaned and not cleaned.startswith('['):
                steps.append(cleaned)

    # Error logs
    error_logs = ""
    logs_match = re.search(r'Error Logs:\s*\n(.*?)(?:\nSuspected Area:|\Z)', block, re.DOTALL)
    if logs_match:
        error_logs = logs_match.group(1).strip()

    # Suspected area
    suspected = {"modules": [], "files": [], "functions": []}
    susp_match = re.search(r'Suspected Area:\s*\n(.*?)(?:\n[─]{10,}|\n={10,}|\Z)', block, re.DOTALL)
    if susp_match:
        susp_text = susp_match.group(1)
        suspected["modules"] = _split_csv(_extract_field(susp_text, "Modules"))
        # Files may span multiple lines
        files_field = _extract_field(susp_text, "Files")
        if files_field:
            suspected["files"] = _split_csv(files_field)
        suspected["functions"] = _split_csv(_extract_field(susp_text, "Functions"))

    return {
        "id": bug_id,
        "title": title,
        "description": description,
        "steps_to_reproduce": steps,
        "expected_behavior": expected,
        "actual_behavior": actual,
        "severity": severity,
        "reported_by": reported_by,
        "environment": environment,
        "error_logs": error_logs,
        "suspected_area": suspected,
    }


def _parse_business_rules(content):
    """Parse SECTION N: BUSINESS RULES."""
    section = _extract_section(content, "SECTION 3: BUSINESS RULES")
    if not section:
        section = _extract_section(content, "SECTION 4: BUSINESS RULES")
    if not section:
        return []

    rules = []
    for line in section.split('\n'):
        line = line.strip()
        # Match "BR-001 : rule text" or "BR-001: rule text"
        match = re.match(r'(BR-\d+)\s*:\s*(.+)', line)
        if match:
            rules.append({"id": match.group(1), "rule": match.group(2).strip()})
    return rules


def _parse_api_contracts(content):
    """Parse SECTION N: HEADLESS API CONTRACT REFERENCE."""
    section = _extract_section(content, "SECTION 4: HEADLESS API CONTRACT REFERENCE")
    if not section:
        section = _extract_section(content, "SECTION 5: HEADLESS API CONTRACT REFERENCE")
    if not section:
        return {"queries": [], "mutations": [], "error_codes": []}

    queries = []
    mutations = []
    error_codes = []

    # Parse GraphQL Queries block
    q_match = re.search(r'GraphQL Queries:\s*\n(.*?)(?:\nGraphQL Mutations:|\n\w)', section, re.DOTALL)
    if q_match:
        for line in q_match.group(1).strip().split('\n'):
            line = line.strip()
            # Match "queryName : description" or "queryName (extended) : desc"
            m = re.match(r'(\S+(?:\s*\([^)]*\))?)\s*:\s*(.+)', line)
            if m:
                queries.append({"name": m.group(1).strip(), "purpose": m.group(2).strip()})

    # Parse GraphQL Mutations block
    m_match = re.search(r'GraphQL Mutations:\s*\n(.*?)(?:\nError Codes:|\n\w[A-Z])', section, re.DOTALL)
    if m_match:
        for line in m_match.group(1).strip().split('\n'):
            line = line.strip()
            m = re.match(r'(\S+)\s*:\s*(.+)', line)
            if m:
                mutations.append({"name": m.group(1).strip(), "purpose": m.group(2).strip()})

    # Parse Error Codes block
    ec_match = re.search(r'Error Codes:\s*\n(.*?)(?:\Z)', section, re.DOTALL)
    if ec_match:
        codes_text = ec_match.group(1).strip()
        # Split by comma or newline
        for code in re.split(r'[,\n]+', codes_text):
            code = code.strip()
            if code and code.isupper():
                error_codes.append(code)

    return {"queries": queries, "mutations": mutations, "error_codes": error_codes}


def _parse_data_model(content):
    """Parse SECTION N: DATA MODEL / PERSISTENCE."""
    section = _extract_section(content, "SECTION 5: DATA MODEL / PERSISTENCE")
    if not section:
        section = _extract_section(content, "SECTION 6: DATA MODEL / PERSISTENCE")
    if not section:
        return []

    tables = []
    # Prepend newline so first entry matches the split pattern
    section = "\n" + section
    # Split by table blocks (table_name:)
    blocks = re.split(r'\n(\w[\w_ ]+(?:extension attributes)?):\s*\n', section)
    # blocks[0] is preamble, then alternating: table_name, content
    for i in range(1, len(blocks) - 1, 2):
        table_name = blocks[i].strip()
        block_content = blocks[i + 1] if i + 1 < len(blocks) else ""
        fields = _extract_field(block_content, "Fields")
        purpose = _extract_field(block_content, "Purpose")
        tables.append({
            "table": table_name,
            "fields": fields,
            "purpose": purpose,
        })

    return tables


def _parse_nfr(content):
    """Parse SECTION N: NON-FUNCTIONAL REQUIREMENTS."""
    section = _extract_section(content, "SECTION 7: NON-FUNCTIONAL REQUIREMENTS")
    if not section:
        section = _extract_section(content, "SECTION 8: NON-FUNCTIONAL REQUIREMENTS")
    if not section:
        return []

    nfrs = []
    for line in section.split('\n'):
        line = line.strip()
        # Match "Category : requirement text"
        m = re.match(r'(\w[\w\s/&]*\w)\s*:\s*(.+)', line)
        if m:
            nfrs.append({"category": m.group(1).strip(), "requirement": m.group(2).strip()})
    return nfrs


def _parse_test_scenarios(content):
    """Parse SECTION N: QA AND TEST SCENARIOS."""
    section = _extract_section(content, "SECTION 8: QA AND TEST SCENARIOS")
    if not section:
        section = _extract_section(content, "SECTION 9: QA AND TEST SCENARIOS")
    if not section:
        return []

    scenarios = []
    for line in section.split('\n'):
        line = line.strip()
        # Match "TC-001 : scenario text"
        m = re.match(r'(TC-\d+)\s*:\s*(.+)', line)
        if m:
            scenarios.append({"id": m.group(1), "scenario": m.group(2).strip()})
    return scenarios
